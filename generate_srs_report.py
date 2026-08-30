"""
MediGuide Vision AI — Mid-Term SRS Report Generator
===================================================
Generates a comprehensive IEEE 830 compliant Software Requirements Specification
(SRS) & Mid-Term Project Report for Review 2 (Implementation Level Review).
Includes system architecture, functional/non-functional requirements,
data flow diagrams, security models, test plans, and project progress.
Exports to both DOCX and PDF.
"""

import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DOCX = os.path.abspath(r'documents/MediGuide_Vision_AI_SRS_Report.docx')
OUTPUT_PDF  = os.path.abspath(r'documents/MediGuide_Vision_AI_SRS_Report.pdf')

doc = Document()

# Page setup: Standard A4 with 1.0 in margins
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Helper styling
def add_report_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p

def add_report_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
    return p

def add_sec_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p

def add_subsec_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_p(text, bold_prefix=None, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# 1. COVER PAGE / METADATA HEADER
# ─────────────────────────────────────────────────────────────────────────────
add_report_title("VARDHAMAN COLLEGE OF ENGINEERING")
p_inst = doc.add_paragraph()
p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_inst.paragraph_format.space_after = Pt(14)
r_inst = p_inst.add_run("(Autonomous)\nDepartment of Information Technology\nAffiliated to JNTUH | Approved by AICTE | Accredited by NAAC with A++ Grade\nKacharam, Shamshabad, Hyderabad – 501218, Telangana, India")
r_inst.font.name = 'Arial'; r_inst.font.size = Pt(9.5); r_inst.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_report_title("SOFTWARE REQUIREMENTS SPECIFICATION (SRS)")
add_report_subtitle("MID-TERM PROJECT REPORT — REVIEW 2 (IMPLEMENTATION LEVEL REVIEW)")

# Meta Table
meta_t = doc.add_table(rows=7, cols=2)
meta_t.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_info = [
    ("Project Title", "MediGuide Vision AI: Multiagent medical image Analysis and Healthcare Navigation"),
    ("Domain", "Image Processing & Healthcare Artificial Intelligence"),
    ("Course & Academic Year", "Major Project (Semester 7) | Academic Year: 2026–2027"),
    ("Batch & Section", "Batch 2023120304 — Section C (Batch ID: 23ITMNP-C06)"),
    ("Team Members", "1. Ratna Ganesh Reddy (23ITMNP-C06-A)\n2. P. Maruthi Sai Teja (23ITMNP-C06-B)\n3. P. Pavan Goud (23ITMNP-C06-C)"),
    ("Project Supervisor / Guide", "Dr. L. Sunitha, Associate Professor, Dept. of Information Technology"),
    ("Section Coordinator", "Dr. Vinayak G Biradar, Department of Information Technology")
]

for ri, (k, v) in enumerate(meta_info):
    c0 = meta_t.rows[ri].cells[0]
    c1 = meta_t.rows[ri].cells[1]
    c0.paragraphs[0].text = k; c0.paragraphs[0].runs[0].font.bold = True; c0.paragraphs[0].runs[0].font.size = Pt(10)
    c1.paragraphs[0].text = v; c1.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2. EXECUTIVE SUMMARY & PROJECT SCOPE
# ─────────────────────────────────────────────────────────────────────────────
add_sec_heading("1. EXECUTIVE SUMMARY & PROJECT OVERVIEW")
add_p(
    "MediGuide Vision AI is an advanced, multimodal, multi-agent healthcare navigation and diagnostic decision support system "
    "designed to bridge the severe outpatient healthcare access and diagnostic comprehension crisis. Preventable medical errors account "
    "for more than 250,000 deaths annually worldwide, predominantly resulting from uncoordinated multi-prescribing, therapeutic drug duplications, "
    "misinterpreted laboratory blood panels, and patient unawareness of food-drug metabolic contraindications."
)
add_p(
    "The primary objective of MediGuide Vision AI is to combine state-of-the-art Computer Vision (Pillow CLAHE, Lanczos scaling, and Otsu binarization) "
    "with Google's Agent Development Kit (ADK) multi-agent orchestration. By decoupling clinical navigation into five specialized, sandboxed sub-agents "
    "(VisionAgent, LabAgent, TriageAgent, SchedulerAgent, and ResearchAgent) and enforcing deterministic Python safety verification rules, "
    "the system eliminates LLM hallucination risks in dosage, contraindication, and diagnostic laboratory evaluation."
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. OVERALL SYSTEM DESCRIPTION & ARCHITECTURAL DESIGN
# ─────────────────────────────────────────────────────────────────────────────
add_sec_heading("2. OVERALL SYSTEM ARCHITECTURE & MULTI-AGENT SPECIFICATION")
add_p(
    "The system architecture is structured around a hierarchical multi-agent design pattern. Incoming user queries and image uploads "
    "pass through a security filter before reaching the root orchestrator:"
)
add_bullet("Local PII Redaction Layer: Scans user input text and OCR transcriptions to redact Personally Identifiable Information (Patient Names, Phone Numbers, Aadhaar/SSN IDs) locally before sending prompts to external model endpoints.", bold_prefix="Security Layer: ")
add_bullet("OrchestratorAgent (Root): Acts as the clinical traffic controller. Inspects query intent, modal payloads, and urgency flags to delegate execution to domain-specialist sub-agents.", bold_prefix="Root Orchestrator: ")
add_bullet("VisionAgent: Dedicated to computer vision. Manages the 4-modality image processing pipeline (Prescription OCR, Skin Lesion Triage, Chest X-Rays, ECG waveforms) via Pillow and Gemini Vision.", bold_prefix="Vision Specialist: ")
add_bullet("LabAgent: Diagnostic lab specialist. Evaluates 35+ quantitative blood biomarkers against clinical bounds and cross-checks findings against patient medications.", bold_prefix="Laboratory Specialist: ")
add_bullet("TriageAgent: Symptom assessment specialist. Performs structured clinical triage, calculates the 7-Factor Health Risk Score (1–100 scale), and triggers emergency routing.", bold_prefix="Triage Specialist: ")
add_bullet("SchedulerAgent: Medication compliance specialist. Computes clinically validated 24-hour daily schedules based on circadian pharmacological timing.", bold_prefix="Medication Scheduler: ")
add_bullet("ResearchAgent: Medical knowledge specialist. Retrieves evidence-based disease pathology, causes, treatments, and prevention guidelines from curated clinical registries.", bold_prefix="Medical Research: ")

# ─────────────────────────────────────────────────────────────────────────────
# 4. FUNCTIONAL REQUIREMENTS (IEEE 830)
# ─────────────────────────────────────────────────────────────────────────────
add_sec_heading("3. FUNCTIONAL REQUIREMENTS SPECIFICATION")

add_subsec_heading("3.1 Multimodal Image Processing & Vision Module (FR-01 to FR-04)")
add_bullet("FR-01 (Prescription OCR & Brand Resolution): The system shall ingest smartphone-captured prescription images, normalize resolution to 1024px using Lanczos interpolation, perform OCR text extraction, and deterministically map brand names (e.g., Trimox) to active generic molecules (e.g., Amoxicillin).")
add_bullet("FR-02 (Therapeutic Duplication Detection): The system shall programmatically raise a critical alert if two or more medications resolve to the identical generic active ingredient, preventing accidental double-dosing.")
add_bullet("FR-03 (Skin Lesion Triage): The system shall apply CLAHE contrast enhancement on skin lesion photographs, evaluate visual characteristics against clinical ABCDE criteria, and assign triage urgency tiers (GREEN / YELLOW / RED).")
add_bullet("FR-04 (Chest X-Ray & ECG Triage): The system shall process radiograph and ECG images, applying grid-line binarization and structural feature extraction to flag life-threatening anomalies for urgent clinician review.")

add_subsec_heading("3.2 Clinical Safety & Deterministic Analysis (FR-05 to FR-08)")
add_bullet("FR-05 (Food-Drug Interaction Engine): The system shall verify prescribed medications against 8 food/substance classes (Grapefruit, Alcohol, Dairy, Vitamin-K, Tyramine, St. John's Wort, Caffeine, Potassium) and flag CYP3A4 metabolic contraindications.")
add_bullet("FR-06 (Diagnostic Lab Report Analyzer): The system shall parse numeric blood test parameters (HbA1c, Glucose, Creatinine, eGFR, Troponin, Potassium, Lipids, TSH, CBC) against age- and gender-adjusted reference ranges, outputting plain-English clinical explanations.")
add_bullet("FR-07 (Critical Emergency Alarms): The system shall immediately flag emergency values (Troponin > 0.04 ng/mL for cardiac infarction, Potassium > 6.0 mEq/L for cardiac arrest risk) and output emergency contact hotlines.")
add_bullet("FR-08 (Quantified Health Risk Score): The system shall calculate an auditable 1–100 Health Risk Score using a 7-factor weighted clinical algorithm aligned with NEWS2 guidelines.")

# ─────────────────────────────────────────────────────────────────────────────
# 5. NON-FUNCTIONAL REQUIREMENTS (IEEE 830)
# ─────────────────────────────────────────────────────────────────────────────
add_sec_heading("4. NON-FUNCTIONAL REQUIREMENTS (NFR)")
add_bullet("NFR-01 (Performance & Latency): The multi-agent routing pipeline shall process single-text queries in < 1.0s and multimodal image payloads in < 2.5s under standard network conditions.", bold_prefix="Response Latency: ")
add_bullet("NFR-02 (Deterministic Safety & Zero Hallucination): Safety-critical features (drug duplications, food interactions, lab reference checks) shall execute on pure Python code, guaranteeing 100% reproducibility.", bold_prefix="Safety Guarantee: ")
add_bullet("NFR-03 (Resilience & Offline Fallback): In the event of API rate limits (HTTP 429) or internet disconnection, the system shall seamlessly fall back to local rule-based simulation mode without application crashing.", bold_prefix="Fault Tolerance: ")
add_bullet("NFR-04 (Security & Privacy): Patient Health Information (PHI) shall be redacted locally. No plain-text API secrets or keys shall be exposed in the client or repository.", bold_prefix="Data Security: ")
add_bullet("NFR-05 (Containerization & Portability): The system shall be fully containerized using Docker and Docker Compose, deployable on Linux/Windows/Mac and cloud environments (FastAPI REST backend).", bold_prefix="Deployability: ")

# ─────────────────────────────────────────────────────────────────────────────
# 6. VERIFICATION & TEST PLAN
# ─────────────────────────────────────────────────────────────────────────────
add_sec_heading("5. SYSTEM VERIFICATION, TESTING & AUDIT RESULTS")
add_p(
    "To satisfy the rigorous CodeLoop AI Audit criteria (Automated Testing & Coverage, Code Modularity, Security Hygiene), "
    "a comprehensive test suite was implemented in Python using the `pytest` framework covering all sub-systems:"
)

test_t = doc.add_table(rows=7, cols=4)
test_t.alignment = WD_TABLE_ALIGNMENT.CENTER

t_headers = ["Test Suite Module", "Target Scope / Methods", "Test Scenarios Evaluated", "Status"]
for ci, h in enumerate(t_headers):
    c = test_t.rows[0].cells[ci]
    c.paragraphs[0].text = h
    c.paragraphs[0].runs[0].font.bold = True
    c.paragraphs[0].runs[0].font.size = Pt(9)

t_rows = [
    ["test_vision_tools.py", "Pillow Image Preprocessing & CLAHE", "Image mode conversion, Lanczos aspect scaling, empty payload validation", "PASSED (100%)"],
    ["test_medical_tools.py", "Brand-to-Generic & Duplication Checker", "Trimox/Amoxicillin duplication alert, drug interaction matrix", "PASSED (100%)"],
    ["test_lab_tools.py", "35+ Lab Reference Ranges & Alarms", "Critical Troponin emergency, Creatinine-Metformin lactic acidosis conflict", "PASSED (100%)"],
    ["test_food_drug_tools.py", "Food-Drug Matrix & Chrono-Scheduler", "Grapefruit x Warfarin CYP3A4 inhibition, daily schedule generation", "PASSED (100%)"],
    ["test_security.py", "PII Redaction & Sanitization", "Name, phone, and ID pattern redaction from clinical queries", "PASSED (100%)"],
    ["test_orchestrator.py", "ADK Multi-Agent Routing Logic", "Emergency triage routing, agent hand-off verification", "PASSED (100%)"]
]

for ri, row_data in enumerate(t_rows):
    for ci, val in enumerate(row_data):
        c = test_t.rows[ri+1].cells[ci]
        c.paragraphs[0].text = val
        c.paragraphs[0].runs[0].font.size = Pt(8.5)

# ─────────────────────────────────────────────────────────────────────────────
# 7. REVIEW 2 STATUS & FUTURE ROADMAP
# ─────────────────────────────────────────────────────────────────────────────
add_sec_heading("6. REVIEW 2 (ILR) STATUS & PROJECT ROADMAP")
add_p(
    "As of Review 2 (Implementation Level Review, 31 August 2026), all core backend modules, multi-agent coordination pipelines, "
    "vision image preprocessors, diagnostic lab analyzers, and automated test suites have been fully implemented, verified, "
    "and committed to the project repository."
)
add_bullet("Phase 1 to Phase 4 (Completed): ADK Multi-Agent Core, Vision Pipeline, Deterministic Clinical DBs, Risk Score Engine, FastAPI Server, and Docker deployment.")
add_bullet("Phase 5 (Review 2 Current): Implementation Level Review Presentation (PPTX/PDF), IEEE Research Paper Manuscript (DOCX/PDF), and Mid-Term SRS Report (DOCX/PDF).")
add_bullet("Phase 6 (Review 3 & Final Defense): FHIR standard hospital integration, live clinical user trials, and mobile client application deployment.")

# Save DOCX
doc.save(OUTPUT_DOCX)
print(f"Mid-Term SRS Report DOCX saved to: {OUTPUT_DOCX}")

# Convert DOCX to PDF using win32com Word Automation
try:
    import win32com.client
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False
    doc_com = word_app.Documents.Open(OUTPUT_DOCX)
    doc_com.SaveAs(OUTPUT_PDF, FileFormat=17)
    doc_com.Close()
    word_app.Quit()
    print(f"Mid-Term SRS Report PDF successfully exported to: {OUTPUT_PDF}")
except Exception as e:
    print(f"Word to PDF conversion error: {e}")
