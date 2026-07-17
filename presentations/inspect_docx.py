import docx

doc_path = r'C:\Users\marut\capstone\presentations\Literature Survey Report.docx'

doc = docx.Document(doc_path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print()

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt:
        print(f"[{i:2d}] \"{txt[:120]}\"")
        for j, run in enumerate(p.runs):
            print(f"    run[{j}]: \"{run.text}\" (font={run.font.name}, sz={run.font.size.pt if run.font.size else 'inh'}, bold={run.font.bold})")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} columns")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"  Row {ri}: {cells}")
