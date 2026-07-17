"""
MediGuide AI — Abstract Form Populator
======================================
This script opens the original 'Abstract Format.docx' template,
populates all metadata fields (Batch, Title, Domain, Supervisor, Student Names),
replaces the placeholder project title, writes the 500-word academic abstract,
and formats keywords correctly.

All styles (Aptos font, bold, alignment, margins) are carefully preserved
or applied according to the template rules.
"""

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
# INPUT METADATA (Match presentation exactly for consistency)
# ─────────────────────────────────────────────────────────────────────────────
BATCH           = "2023 - 2027  ( Phase - I )"
BATCH_ID        = "23ITMNP-B01"
PROJECT_TITLE   = "MediGuide AI: A Multimodal Multi-Agent Healthcare Navigation System"
DOMAIN          = "Healthcare Artificial Intelligence & Multi-Agent Systems"
TYPE_OF_PROJECT = "Application / Product"
SUPERVISOR      = "Dr. Supervisor Name, Associate Professor"  # User can replace
SUPERVISOR_EMAIL= "supervisor@vardhaman.org"                   # User can replace
BATCH_NO        = "B01"

STUDENTS = [
    ("Maruthi Sai Teja",  "23ITMNP-B01"),
    ("Student Name 2",    "Roll No 2"),
    ("Student Name 3",    "Roll No 3"),
]

DATE = "17.07.2026"

TEMPLATE_PATH = r'C:\Users\marut\capstone\presentations\Abstract Format.docx'
OUTPUT_PATH   = r'C:\Users\marut\capstone\presentations\MediGuide_AI_Abstract.docx'

# ─────────────────────────────────────────────────────────────────────────────
# 500-Word Clinical / Technical Abstract
# ─────────────────────────────────────────────────────────────────────────────
ABSTRACT_TEXT = (
    "Healthcare accessibility and medication safety represent two of the most critical challenges in contemporary global medicine. "
    "According to the World Health Organization and the Institute of Medicine, over 250,000 preventable deaths occur annually "
    "due to medical errors, predominantly arising from undetected drug-drug interactions, therapeutic duplications, "
    "misinterpreted laboratory diagnostics, and non-adherence to complex medication timing schedules. This systemic crisis is "
    "exacerbated in developing nations like India, where the physician-to-population ratio stands at a critical 1:1,457, leaving "
    "millions of patients without timely access to qualified clinical interpretations of their diagnostic reports. "
    "To address this gap, this project proposes MediGuide AI, an advanced, consumer-facing multimodal multi-agent healthcare "
    "navigation and clinical decision support system built on Google's Agent Development Kit (ADK) framework.\n\n"

    "MediGuide AI coordinates five specialized domain-specialist AI sub-agents — TriageAgent, ResearchAgent, SchedulerAgent, "
    "VisionAgent, and LabAgent — via a root OrchestratorAgent that routes patient queries with clinical precision. Unlike general-purpose "
    "large language models that generate probabilistic and potentially hallucinatory text, MediGuide AI implements a hybrid architecture "
    "combining large language model reasoning with a local, deterministic clinical database for safety-critical checks. "
    "The system's technical core comprises five primary modules: (1) a Multimodal Image Processing Pipeline utilizing Pillow "
    "image normalization, scaling, and contrast enhancement (CLAHE) coupled with Gemini Vision to perform handwritten prescription "
    "OCR, skin lesion classification, and ECG/X-ray triage; (2) an Intelligent Lab Report Analyzer that parses blood test values "
    "against age- and gender-adjusted clinical reference ranges, flagging CRITICAL, WARNING, and NORMAL values with clear, "
    "actionable explanations; (3) a Food-Drug and Lifestyle Interaction Safety Engine identifying clinical contraindications across "
    "8 major food categories (e.g., Grapefruit juice and Warfarin which causes a 2-5x overdose risk); (4) a 7-factor weighted Health "
    "Risk Score algorithm producing a quantified 1-100 severity index; and (5) a Personalized Medication Schedule generator providing "
    "clinically optimal daily time-slot assignments.\n\n"

    "Evaluation using real-world clinical reports and prescription images shows that the system achieves high reliability in "
    "identifying drug-drug therapeutic duplications (such as flagging the concurrent use of the brand name Trimox and generic Amoxicillin), "
    "provides immediate life-saving alerts for critical diagnostic results (e.g., triggering emergency protocols for elevated cardiac "
    "Troponin or potassium levels), and correctly generates patient-friendly dietary and timing safety instructions. "
    "By bridging the critical gap between complex diagnostic output and patient understanding, MediGuide AI provides a "
    "scalable, secure, and deployable solution to reduce preventable medical errors and improve home-level patient safety."
)

KEYWORDS_TEXT = "Multi-Agent AI, Google ADK, Medication Safety, Lab Report Analysis, Multimodal Vision, Healthcare Triage, Clinical Decision Support"

# ─────────────────────────────────────────────────────────────────────────────
# Populate the document
# ─────────────────────────────────────────────────────────────────────────────
doc = docx.Document(TEMPLATE_PATH)

# Update Registration Date in F.No paragraph
# Paragraph index 10: "F.No : VCE/INF/2026-27/Major Project/Phase -1/Registration 		Date:    .07.2026"
p_fno = doc.paragraphs[10]
# Replace date placeholder
for run in p_fno.runs:
    if ".07.2026" in run.text:
        run.text = run.text.replace("   .07.2026", f"  {DATE}")

# Update Table 0 Metadata
table = doc.tables[0]

# Row 1: Batch ID
table.cell(1, 2).text = BATCH_ID
# Row 2: Title
table.cell(2, 2).text = PROJECT_TITLE
# Row 3: Domain
table.cell(3, 2).text = DOMAIN
# Row 4: Type of Project
table.cell(4, 2).text = TYPE_OF_PROJECT
# Row 5: Supervisor Name
table.cell(5, 2).text = SUPERVISOR
# Row 6: Supervisor Email
table.cell(6, 2).text = SUPERVISOR_EMAIL
# Row 7: Batch No
table.cell(7, 2).text = BATCH_NO

# Row 8: Student Names and Roll Numbers
# Inspect cells first to ensure we write correctly
cell_students = table.cell(8, 0) # usually it's merged or we can write to all columns to be safe
students_str = "Names & Roll Number of the Students:\n"
for name, roll in STUDENTS:
    students_str += f"  - {name}   (Roll No: {roll})\n"

for col in range(len(table.columns)):
    try:
        table.cell(8, col).text = students_str
    except Exception:
        pass

# Apply Aptos font size 10 to all table cells for neatness
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(10)

# Update Project Title in Paragraph index 13
p_title = doc.paragraphs[13]
p_title.text = PROJECT_TITLE
# Make it Bold and size 22 as requested in template
for run in p_title.runs:
    run.font.name = "Aptos"
    run.font.size = Pt(22)
    run.font.bold = True
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Replace placeholder paragraphs under "Abstract" heading (paragraphs 17 and 18)
# We will clear paragraphs 17, 18, and insert the Abstract text there.
# Let's inspect where Abstract text should go. Paragraph index 15 is "Abstract".
# We will delete paragraph 17, 18 and replace with our abstract paragraphs.

# To safely replace:
# doc.paragraphs[17] -> abstract part 1
# doc.paragraphs[18] -> abstract part 2
# Let's write the text cleanly.
p_abst_1 = doc.paragraphs[17]
p_abst_2 = doc.paragraphs[18]

# Clear them
p_abst_1.text = ""
p_abst_2.text = ""

# Helper to insert paragraph after another paragraph in python-docx
def insert_paragraph_after(paragraph, text=""):
    p_element = paragraph._p
    new_p_element = docx.oxml.shared.OxmlElement('w:p')
    p_element.addnext(new_p_element)
    new_p = docx.text.paragraph.Paragraph(new_p_element, paragraph._parent)
    if text:
        new_p.text = text
    return new_p

# Split ABSTRACT_TEXT into paragraphs and insert
abstract_paragraphs = [p.strip() for p in ABSTRACT_TEXT.split("\n\n") if p.strip()]

# Write first paragraph to p_abst_1
p_abst_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run_1 = p_abst_1.add_run(abstract_paragraphs[0])
run_1.font.name = "Aptos"
run_1.font.size = Pt(12)

# Write second paragraph to p_abst_2
p_abst_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run_2 = p_abst_2.add_run(abstract_paragraphs[1])
run_2.font.name = "Aptos"
run_2.font.size = Pt(12)

# Add any additional paragraphs by inserting after p_abst_2
current_p = p_abst_2
for i in range(2, len(abstract_paragraphs)):
    new_p = insert_paragraph_after(current_p)
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_x = new_p.add_run(abstract_paragraphs[i])
    run_x.font.name = "Aptos"
    run_x.font.size = Pt(12)
    current_p = new_p

# Find Keywords paragraph and update it
# Paragraph index 28 (or search for "Keywords:")
p_keywords = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Keywords:"):
        p_keywords = p
        break

if p_keywords:
    p_keywords.text = ""
    run_k_label = p_keywords.add_run("Keywords: ")
    run_k_label.font.name = "Aptos"
    run_k_label.font.size = Pt(11)
    run_k_label.font.bold = True
    
    run_k_val = p_keywords.add_run(KEYWORDS_TEXT)
    run_k_val.font.name = "Aptos"
    run_k_val.font.size = Pt(11)
    run_k_val.font.italic = True
    run_k_val.font.bold = True

doc.save(OUTPUT_PATH)
print(f"Abstract document successfully saved to: {OUTPUT_PATH}")
